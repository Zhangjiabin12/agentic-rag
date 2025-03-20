import os
import re
import json
import logging
import argparse
import shutil
import sys
import time
import glob
from typing import List, Dict, Optional, Any
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor, as_completed

from config import config

import pandas as pd
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
import pypdfium2
import chromadb

from tools.rag_tools.text_splitter import TextSplitter
from tools.rag_tools.text_embedder import TextEmbedder
from tools.rag_tools.vector_retriever import VectorRetriever, VectorDB

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger('DocumentProcessor')

class Document:
    """文档类，用于存储文档内容和元数据"""
    
    def __init__(self, page_content: str, metadata: Optional[Dict] = None):
        """初始化文档
        
        Args:
            page_content: 文档内容
            metadata: 文档元数据
        """
        self.page_content = page_content
        self.metadata = metadata or {}

class BaseExtractor:
    """提取器基类"""
    
    def extract(self) -> List[Document]:
        """提取文档
        
        Returns:
            文档列表
        """
        raise NotImplementedError("子类必须实现extract方法")

class TextExtractor:
    """简化版文本提取器"""
    
    def __init__(self):
        """初始化文本提取器"""
        pass
    
    def extract(self, file_path: str) -> Dict[str, Any]:
        """提取文本并转换为UTF-8编码
        
        Args:
            file_path: 文件路径
            
        Returns:
            提取结果，包括文本和元数据
        """
        logger.info(f"提取文本: {file_path}")
        
        # 获取文件扩展名
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()
        
        # 如果是PDF文件，使用PdfExtractor处理
        if ext == '.pdf':
            try:
                pdf_extractor = PdfExtractor(file_path)
                documents = pdf_extractor.extract()
                
                if not documents:
                    raise RuntimeError("PDF文件未提取到任何内容")
                
                # 合并所有页面的内容
                text = "\n\n".join(doc.page_content for doc in documents)
                metadata = {
                    "source": file_path,
                    "original_encoding": "utf-8",
                    "converted_to_utf8": False,
                    "total_pages": len(documents)
                }
                
                return {
                    "text": text,
                    "metadata": metadata
                }
            except Exception as e:
                logger.error(f"PDF提取失败: {file_path}, 错误: {str(e)}")
                raise RuntimeError(f"PDF提取失败: {file_path}") from e
        
        # 对于其他文本文件，使用二进制模式读取并检测编码
        try:
            # 读取文件的二进制内容
            with open(file_path, 'rb') as f:
                raw_data = f.read()
            
            # 检测编码
            import chardet
            result = chardet.detect(raw_data)
            detected_encoding = result["encoding"]
            confidence = result["confidence"]
            
            logger.info(f"检测到编码: {detected_encoding} (置信度: {confidence:.2f})")
            
            # 如果检测到的编码可信度较低，或者检测失败，尝试常用编码
            if not detected_encoding or confidence < 0.8:
                encodings_to_try = ['utf-8', 'gbk', 'gb2312', 'gb18030', 'big5', 'utf-16', 'utf-32']
                for encoding in encodings_to_try:
                    try:
                        text = raw_data.decode(encoding)
                        detected_encoding = encoding
                        logger.info(f"使用 {encoding} 编码成功解码文件")
                        break
                    except UnicodeDecodeError:
                        continue
                else:
                    raise RuntimeError(f"无法使用任何已知编码解码文件: {file_path}")
            else:
                # 使用检测到的编码解码
                try:
                    text = raw_data.decode(detected_encoding)
                except UnicodeDecodeError:
                    # 如果解码失败，尝试使用 errors='replace' 参数
                    text = raw_data.decode(detected_encoding, errors='replace')
                    logger.warning(f"使用 {detected_encoding} 编码解码时出现错误，已使用替换字符")
            
            # 记录原始编码和转换信息
            metadata = {
                "source": file_path,
                "original_encoding": detected_encoding,
                "converted_to_utf8": detected_encoding.lower() != 'utf-8',
                "decode_confidence": confidence
            }
            
            return {
                "text": text,
                "metadata": metadata
            }
            
        except Exception as e:
            logger.error(f"读取文件失败: {file_path}, 错误类型: {type(e).__name__}, 错误信息: {str(e)}")
            raise RuntimeError(f"读取文件失败: {file_path}") from e

class MarkdownExtractor(BaseExtractor):
    """Markdown提取器"""
    
    def __init__(self, file_path: str, remove_hyperlinks: bool = False, remove_images: bool = False, encoding: Optional[str] = None, autodetect_encoding: bool = True):
        """初始化Markdown提取器
        
        Args:
            file_path: 文件路径
            remove_hyperlinks: 是否移除超链接
            remove_images: 是否移除图片
            encoding: 编码
            autodetect_encoding: 是否自动检测编码
        """
        self.file_path = file_path
        self.remove_hyperlinks = remove_hyperlinks
        self.remove_images = remove_images
        self.encoding = encoding
        self.autodetect_encoding = autodetect_encoding
        self.text_extractor = TextExtractor()
    
    def parse_tups(self, file_path: str) -> List[tuple]:
        """解析Markdown文件为元组列表
        
        Args:
            file_path: 文件路径
            
        Returns:
            元组列表，每个元组包含标题和内容
        """
        # 读取文件
        text_docs = self.text_extractor.extract(file_path)
        if not text_docs:
            return []
        
        markdown_text = text_docs['text']
        
        # 处理超链接和图片
        if self.remove_hyperlinks:
            markdown_text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', markdown_text)
        
        if self.remove_images:
            markdown_text = re.sub(r'!\[[^\]]*\]\([^)]+\)', '', markdown_text)
        
        # 分割标题和内容
        headers = []
        current_header = None
        current_content = []
        
        for line in markdown_text.split('\n'):
            header_match = re.match(r'^(#+)\s+(.*)', line)
            if header_match:
                # 如果有当前标题，保存它和它的内容
                if current_header is not None:
                    headers.append((current_header, '\n'.join(current_content)))
                
                # 开始新的标题
                current_header = line
                current_content = []
            else:
                current_content.append(line)
        
        # 添加最后一个标题和内容
        if current_header is not None:
            headers.append((current_header, '\n'.join(current_content)))
        elif current_content:
            # 如果没有标题但有内容，使用None作为标题
            headers.append((None, '\n'.join(current_content)))
        
        return headers
    
    def extract(self) -> List[Document]:
        """提取文档
        
        Returns:
            文档列表
        """
        tups = self.parse_tups(self.file_path)
        documents = []
        
        for header, value in tups:
            value = value.strip()
            if header is None:
                documents.append(Document(page_content=value, metadata={"source": self.file_path}))
            else:
                documents.append(Document(page_content=f"\n\n{header}\n{value}", metadata={"source": self.file_path, "header": header}))
        
        return documents

class PdfExtractor(BaseExtractor):
    """PDF提取器"""
    
    def __init__(self, file_path: str):
        """初始化PDF提取器
        
        Args:
            file_path: 文件路径
        """
        self.file_path = file_path
    
    def extract(self) -> List[Document]:
        """提取文档
        
        Returns:
            文档列表
        """
        documents = []
        
        try:
            # 将文件路径转换为绝对路径，并确保使用正确的路径分隔符
            abs_path = os.path.abspath(self.file_path)
            abs_path = os.path.normpath(abs_path)
            
            logger.info(f"开始读取PDF文件: {abs_path}")
            
            # 使用二进制模式读取文件内容
            with open(abs_path, 'rb') as file:
                pdf_data = file.read()
            
            # 从内存中加载PDF
            pdf_reader = pypdfium2.PdfDocument(pdf_data, autoclose=True)
            
            total_pages = len(pdf_reader)
            logger.info(f"PDF文件共有 {total_pages} 页")
            
            for page_number, page in enumerate(pdf_reader, start=1):
                try:
                    text_page = page.get_textpage()
                    content = text_page.get_text_bounded()
                    
                    # 清理文本内容
                    content = content.strip()
                    # 移除重复的空白字符
                    content = re.sub(r'\s+', ' ', content)
                    # 移除零宽字符
                    content = re.sub(r'[\u200b\u200c\u200d\ufeff]', '', content)
                    
                    if content:  # 只添加非空内容
                        metadata = {
                            "source": self.file_path,
                            "page": page_number,
                            "total_pages": total_pages
                        }
                        documents.append(Document(page_content=content, metadata=metadata))
                    
                    text_page.close()
                except Exception as e:
                    logger.warning(f"处理第 {page_number} 页时出错: {str(e)}")
                finally:
                    page.close()
            
            pdf_reader.close()
            
            if not documents:
                logger.warning(f"PDF文件 {self.file_path} 未提取到任何文本内容")
            else:
                logger.info(f"成功从PDF文件中提取了 {len(documents)} 页有效内容")
            
            return documents
            
        except FileNotFoundError:
            logger.error(f"PDF文件不存在: {self.file_path}")
            raise RuntimeError(f"PDF文件不存在: {self.file_path}")
        except PermissionError:
            logger.error(f"没有权限读取PDF文件: {self.file_path}")
            raise RuntimeError(f"没有权限读取PDF文件: {self.file_path}")
        except Exception as e:
            logger.error(f"读取PDF文件失败: {self.file_path}, 错误类型: {type(e).__name__}, 错误信息: {str(e)}")
            raise RuntimeError(f"读取PDF文件失败: {self.file_path}, 错误: {str(e)}") from e

class HtmlExtractor(BaseExtractor):
    """HTML提取器"""
    
    def __init__(self, file_path: str):
        """初始化HTML提取器
        
        Args:
            file_path: 文件路径
        """
        self.file_path = file_path
    
    def extract(self) -> List[Document]:
        """提取文档
        
        Returns:
            文档列表
        """
        try:
            with open(self.file_path, 'rb') as f:
                soup = BeautifulSoup(f, 'html.parser')
                text = soup.get_text()
                text = text.strip() if text else ""
            
            metadata = {"source": self.file_path}
            return [Document(page_content=text, metadata=metadata)]
        except Exception as e:
            logger.error(f"读取HTML文件失败: {self.file_path}, 错误: {e}")
            raise RuntimeError(f"读取HTML文件失败: {self.file_path}") from e

class ExcelExtractor(BaseExtractor):
    """Excel提取器"""
    
    def __init__(self, file_path: str):
        """初始化Excel提取器
        
        Args:
            file_path: 文件路径
        """
        self.file_path = file_path
    
    def extract(self) -> List[Document]:
        """提取文档
        
        Returns:
            文档列表
        """
        documents = []
        file_extension = os.path.splitext(self.file_path)[-1].lower()
        
        try:
            if file_extension == ".xlsx":
                wb = pd.ExcelFile(self.file_path, engine="xlrd")
                for sheet_name in wb.sheet_names:
                    df = wb.parse(sheet_name=sheet_name)
                    df.dropna(how='all', inplace=True)
                    
                    for index, row in df.iterrows():
                        page_content = []
                        for k, v in row.items():
                            if pd.notna(v):
                                page_content.append(f'"{k}":"{v}"')
                        
                        metadata = {"source": self.file_path, "sheet": sheet_name, "row": index + 2}
                        documents.append(Document(page_content=";".join(page_content), metadata=metadata))
            
            elif file_extension == ".xls":
                excel_file = pd.ExcelFile(self.file_path, engine="xlrd")
                for sheet_name in excel_file.sheet_names:
                    df = excel_file.parse(sheet_name=sheet_name)
                    df.dropna(how='all', inplace=True)
                    
                    for index, row in df.iterrows():
                        page_content = []
                        for k, v in row.items():
                            if pd.notna(v):
                                page_content.append(f'"{k}":"{v}"')
                        
                        metadata = {"source": self.file_path, "sheet": sheet_name, "row": index + 2}
                        documents.append(Document(page_content=";".join(page_content), metadata=metadata))
            else:
                raise ValueError(f"不支持的文件扩展名: {file_extension}")
        except Exception as e:
            logger.error(f"读取Excel文件失败: {self.file_path}, 错误: {e}")
            raise RuntimeError(f"读取Excel文件失败: {self.file_path}") from e
        
        return documents

class WordExtractor(BaseExtractor):
    """Word提取器"""
    
    def __init__(self, file_path: str):
        """初始化Word提取器
        
        Args:
            file_path: 文件路径
        """
        self.file_path = file_path
    
    def extract(self) -> List[Document]:
        """提取文档
        
        Returns:
            文档列表
        """
        try:
            doc = DocxDocument(self.file_path)
            
            # 提取文本
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            
            # 提取表格
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text)
                    table_text.append(" | ".join(row_text))
                full_text.append("\n".join(table_text))
            
            content = "\n\n".join(full_text)
            metadata = {"source": self.file_path}
            
            return [Document(page_content=content, metadata=metadata)]
        except Exception as e:
            logger.error(f"读取Word文件失败: {self.file_path}, 错误: {e}")
            raise RuntimeError(f"读取Word文件失败: {self.file_path}") from e

class CSVExtractor:
    """简化版CSV提取器"""
    
    def __init__(self):
        """初始化CSV提取器"""
        pass
    
    def extract(self, file_path: str) -> Dict[str, Any]:
        """提取CSV文件并转换为UTF-8编码
        
        Args:
            file_path: 文件路径
            
        Returns:
            提取结果，包括文本和元数据
        """
        logger.info(f"提取CSV: {file_path}")
        
        text = ""
        original_encoding = None
        
        try:
            # 首先尝试使用UTF-8编码读取
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            original_encoding = 'utf-8'
            logger.info(f"CSV文件 {file_path} 使用UTF-8编码读取成功")
        except UnicodeDecodeError:
            # 尝试检测编码
            import chardet
            
            # 读取文件的前4096个字节来检测编码
            with open(file_path, 'rb') as f:
                raw_data = f.read(4096)
            
            result = chardet.detect(raw_data)
            detected_encoding = result["encoding"]
            
            try:
                with open(file_path, 'r', encoding=detected_encoding) as f:
                    text = f.read()
                original_encoding = detected_encoding
                logger.info(f"CSV文件 {file_path} 使用检测到的编码 {detected_encoding} 读取成功")
            except Exception as e:
                logger.error(f"读取CSV文件失败: {file_path}, 错误: {e}")
                raise RuntimeError(f"读取CSV文件失败: {file_path}") from e
        
        # 记录原始编码和转换信息
        metadata = {
            "source": file_path,
            "original_encoding": original_encoding,
            "converted_to_utf8": True if original_encoding.lower() != 'utf-8' else False,
            "file_type": "csv"
        }
        
        # 如果原始编码不是UTF-8，记录转换信息
        if original_encoding and original_encoding.lower() != 'utf-8':
            logger.info(f"CSV文件 {file_path} 从 {original_encoding} 转换为 UTF-8")
        
        return {
            "text": text,
            "metadata": metadata
        }

class DocumentProcessor:
    """文档处理器，用于清洗和分割文档"""
    
    @staticmethod
    def normalize_collection_name(name: str) -> str:
        """规范化集合名称以符合ChromaDB的要求:
        1. 长度在3-63个字符之间
        2. 以字母数字字符开头和结尾
        3. 只能包含字母数字字符、下划线或连字符(-)
        4. 不能包含两个连续的点(..)
        5. 不能是有效的IPv4地址
        """
        import re
        import pinyin
        
        # 如果名称全是ASCII字符，只需要替换非法字符
        if all(ord(c) < 128 for c in name):
            # 替换非法字符为下划线
            normalized = re.sub(r'[^a-zA-Z0-9-]', '_', name)
        else:
            # 对于包含中文的名称，转换为拼音
            normalized = pinyin.get(name, format='strip', delimiter='_')
            # 替换任何非法字符为下划线
            normalized = re.sub(r'[^a-zA-Z0-9-]', '_', normalized)
        
        # 移除开头和结尾的下划线
        normalized = normalized.strip('_')
        
        # 确保以字母开头（如果以数字或其他字符开头，添加前缀）
        if not normalized[0].isalpha():
            normalized = 'col_' + normalized
        
        # 如果长度小于3，添加填充
        if len(normalized) < 3:
            normalized = normalized + '_col'
        
        # 如果长度超过62，截断
        if len(normalized) > 62:
            normalized = normalized[:62]
            # 确保截断后不以下划线结尾
            normalized = normalized.rstrip('_')
        
        # 确保以字母或数字结尾
        if not normalized[-1].isalnum():
            normalized = normalized + 'x'
        
        # 移除连续的下划线
        normalized = re.sub(r'_+', '_', normalized)
        
        return normalized 
    
    def __init__(self, collection_name: str = None):
        """初始化文档处理器
        
        Args:
            collection_name: 集合名称
        """
        # 使用默认集合名称
        self.collection_name = collection_name or config.COLLECTION_NAME
        
        # 规范化集合名称
        self.original_name = self.collection_name
        self.normalized_name = self.normalize_collection_name(self.collection_name)
        
        logger.info(f"初始化DocumentProcessor，原始集合名称: {self.original_name}")
        logger.info(f"规范化后的集合名称: {self.normalized_name}")
        
        # 设置目录路径
        self.dataset_dir = os.path.join("./data/tmp/dataset", self.normalized_name)
        self.split_dir = os.path.join("./data/tmp/split", self.normalized_name)
        
        logger.info(f"清洗目录: {self.dataset_dir}")
        logger.info(f"分块目录: {self.split_dir}")
        
        # 创建必要的目录
        os.makedirs(self.dataset_dir, exist_ok=True)
        os.makedirs(self.split_dir, exist_ok=True)
        
        # 设置分块参数
        self.chunk_size = config.CHUNK_SIZE
        self.chunk_overlap = config.CHUNK_OVERLAP
        self.split_method = config.SPLIT_METHOD
        
        logger.info(f"分块大小: {self.chunk_size}, 重叠大小: {self.chunk_overlap}")
        logger.info(f"分块方法: {self.split_method}")
        
        # 设置最大工作线程数
        self.max_workers = min(4, (os.cpu_count() or 1))
        logger.info(f"最大工作线程数: {self.max_workers}")
        
        # 初始化文本嵌入器
        self.embedder = TextEmbedder(
            model_name=config.DEFAULT_EMBEDDING_MODEL,
            cache_dir=config.EMBEDDING_CACHE_PATH,
            db_path=config.VECTOR_DB_PATH
        )
        
        # 初始化向量数据库
        self.vector_db = VectorDB(db_path=config.VECTOR_DB_PATH)
        
        # 初始化文本提取器和分块器
        self.text_extractor = TextExtractor()
        self.text_splitter = TextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap
        )

    def vectorize_chunks(self, chunk_files: List[str], delete_existing: bool = False) -> Dict:
        """向量化所有分块
        
        Args:
            chunk_files: 分块文件列表
            delete_existing: 是否删除现有集合
            
        Returns:
            处理结果
        """
        logger.info(f"原始集合名称: {self.original_name}")
        logger.info(f"规范化后的集合名称: {self.normalized_name}")
        
        try:
            # 读取所有分块
            documents = []
            metadatas = []
            
            for chunk_file in chunk_files:
                # 读取分块内容
                with open(chunk_file, 'r', encoding='utf-8') as f:
                    text = f.read()
                
                # 读取元数据
                metadata_file = chunk_file.replace('.txt', '.json')
                if os.path.exists(metadata_file):
                    with open(metadata_file, 'r', encoding='utf-8') as f:
                        metadata = json.load(f)
                else:
                    metadata = {
                        "source": os.path.basename(chunk_file),
                        "chunk_index": len(documents)
                    }
                
                documents.append(text)
                metadatas.append(metadata)
            
            # 获取集合特定的数据库路径
            db_path = os.path.join(config.VECTOR_DB_PATH, self.normalized_name)
            os.makedirs(db_path, exist_ok=True)
            
            # 使用集合特定的客户端
            client = chromadb.PersistentClient(path=db_path)
            
            # 如果需要删除现有集合
            if delete_existing:
                try:
                    client.delete_collection(name=self.normalized_name)
                    logger.info(f"已删除现有集合: {self.normalized_name}")
                except Exception as e:
                    logger.warning(f"删除集合失败（可能不存在）: {e}")
            
            try:
                # 尝试获取现有集合
                collection = client.get_collection(name=self.normalized_name)
                logger.info(f"使用现有集合: {self.normalized_name}")
            except Exception:
                # 如果集合不存在，创建新集合
                collection = client.create_collection(
                    name=self.normalized_name,
                    metadata={
                        "original_name": self.original_name,
                        "embedding_model": config.DEFAULT_EMBEDDING_MODEL
                    }
                )
                logger.info(f"创建新集合: {self.normalized_name}")
            
            # 生成唯一ID
            timestamp = int(time.time())
            ids = [f"{self.normalized_name}_{timestamp}_{i:06d}" for i in range(len(documents))]
            
            # 批量添加文档
            self.embedder.add_documents(
                collection_name=self.normalized_name,
                documents=documents,
                metadatas=metadatas,
                ids=ids
            )
            
            logger.info(f"向量化完成，共处理 {len(documents)} 个分块")
            return {
                "collection_name": self.normalized_name,
                "chunk_count": len(documents),
                "message": f"成功向量化 {len(documents)} 个分块"
            }
            
        except Exception as e:
            error_msg = f"向量化失败: {str(e)}"
            logger.error(error_msg)
            raise RuntimeError(error_msg)

    def process_file(self, file_path: str, delete_existing: bool = False) -> Dict[str, Any]:
        """处理单个文件，包括清洗、分块和向量化
        
        Args:
            file_path: 文件路径
            delete_existing: 是否删除已存在的集合
            
        Returns:
            处理结果，包括清洗、分块和向量化的信息
        """
        logger.info(f"开始处理文件: {file_path}")
        
        # 检查集合是否存在
        if delete_existing and self.collection_exists(self.collection_name):
            logger.info(f"集合已存在: {self.original_name}，正在删除...")
            success = self.delete_collection(self.collection_name)
            if not success:
                raise RuntimeError(f"删除集合失败: {self.original_name}")
            logger.info(f"集合已删除: {self.original_name}")
        
        # 获取文件扩展名
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()
        
        # 根据文件类型选择合适的提取器
        try:
            if ext == '.csv':
                clean_result = self.csv_extractor.extract(file_path)
            elif ext == '.docx':
                word_extractor = WordExtractor(file_path)
                documents = word_extractor.extract()
                if not documents:
                    raise RuntimeError(f"Word文档未提取到任何内容: {file_path}")
                
                # 合并所有文档内容
                text = "\n\n".join(doc.page_content for doc in documents)
                metadata = documents[0].metadata
                clean_result = {
                    "text": text,
                    "metadata": metadata
                }
                logger.info(f"成功从Word文档中提取了 {len(text)} 个字符")
            elif ext == '.pdf':
                pdf_extractor = PdfExtractor(file_path)
                documents = pdf_extractor.extract()
                if not documents:
                    raise RuntimeError(f"PDF文件未提取到任何内容: {file_path}")
                
                # 合并所有页面的内容
                text = "\n\n".join(doc.page_content for doc in documents)
                metadata = {
                    "source": file_path,
                    "total_pages": len(documents),
                    "file_type": "pdf"
                }
                clean_result = {
                    "text": text,
                    "metadata": metadata
                }
                logger.info(f"成功从PDF文件中提取了 {len(documents)} 页，共 {len(text)} 个字符")
            elif ext == '.md':
                markdown_extractor = MarkdownExtractor(
                    file_path,
                    remove_hyperlinks=True,  # 移除超链接，只保留文本
                    remove_images=True,      # 移除图片引用
                    autodetect_encoding=True # 自动检测编码
                )
                documents = markdown_extractor.extract()
                if not documents:
                    raise RuntimeError(f"Markdown文件未提取到任何内容: {file_path}")
                
                # 合并所有文档内容
                text = "\n\n".join(doc.page_content for doc in documents)
                metadata = {
                    "source": file_path,
                    "file_type": "markdown",
                    "sections": len(documents)
                }
                clean_result = {
                    "text": text,
                    "metadata": metadata
                }
                logger.info(f"成功从Markdown文件中提取了 {len(documents)} 个部分，共 {len(text)} 个字符")
            elif ext == '.html':
                html_extractor = HtmlExtractor(file_path)
                documents = html_extractor.extract()
                if not documents:
                    raise RuntimeError(f"HTML文件未提取到任何内容: {file_path}")
                
                # 合并所有文档内容
                text = "\n\n".join(doc.page_content for doc in documents)
                metadata = {
                    "source": file_path,
                    "file_type": "html"
                }
                clean_result = {
                    "text": text,
                    "metadata": metadata
                }
                logger.info(f"成功从HTML文件中提取了 {len(text)} 个字符")
            elif ext in ['.xlsx', '.xls']:
                excel_extractor = ExcelExtractor(file_path)
                documents = excel_extractor.extract()
                if not documents:
                    raise RuntimeError(f"Excel文件未提取到任何内容: {file_path}")
                
                # 合并所有单元格内容
                text = "\n\n".join(doc.page_content for doc in documents)
                metadata = {
                    "source": file_path,
                    "file_type": "excel",
                    "rows": len(documents)
                }
                clean_result = {
                    "text": text,
                    "metadata": metadata
                }
                logger.info(f"成功从Excel文件中提取了 {len(documents)} 行数据")
            else:
                # 对于其他文本文件，使用通用文本提取器
                clean_result = self.text_extractor.extract(file_path)
                logger.info(f"使用通用文本提取器处理文件: {file_path}")
        except Exception as e:
            logger.error(f"文件提取失败: {file_path}, 错误类型: {type(e).__name__}, 错误信息: {str(e)}")
            raise RuntimeError(f"文件提取失败: {file_path}, 错误: {str(e)}") from e
        
        # 获取清洗后的文本和元数据
        text = clean_result['text']
        metadata = clean_result['metadata']
        
        # 在元数据中添加集合信息
        metadata["collection"] = self.original_name
        metadata["normalized_collection"] = self.normalized_name
        
        # 构建清洗后的文件路径
        clean_file_name = os.path.basename(file_path)
        clean_file_path = os.path.join(self.dataset_dir, clean_file_name)
        
        # 保存清洗后的文本（使用UTF-8编码）
        with open(clean_file_path, 'w', encoding='utf-8') as f:
            f.write(text)
        
        # 保存元数据（使用UTF-8编码）
        metadata_file_path = clean_file_path.replace(ext, '.json')
        with open(metadata_file_path, 'w', encoding='utf-8') as f:
            json.dump(metadata, f, ensure_ascii=False, indent=2)
        
        logger.info(f"文件清洗完成，保存到: {clean_file_path}")
        
        # 分割文本
        chunk_files = self.text_splitter.split_file(
            clean_file_path, 
            output_dir=self.split_dir,
            split_method=self.split_method
        )
        
        # 记录分块信息
        chunk_sizes = []
        for chunk_file in chunk_files:
            try:
                with open(chunk_file, 'r', encoding='utf-8') as f:
                    chunk_text = f.read()
                    chunk_sizes.append(len(chunk_text))
            except Exception as e:
                logger.error(f"读取分块文件失败: {chunk_file}, 错误: {e}")
        
        if chunk_sizes:
            avg_chunk_size = sum(chunk_sizes) / len(chunk_sizes)
            max_chunk_size = max(chunk_sizes)
            min_chunk_size = min(chunk_sizes)
            logger.info(f"分块完成，共 {len(chunk_files)} 个分块")
            logger.info(f"分块大小统计 - 平均: {avg_chunk_size:.2f}, 最大: {max_chunk_size}, 最小: {min_chunk_size}")
            
            # 检查是否有超过设定大小的分块
            oversized_chunks = [size for size in chunk_sizes if size > self.chunk_size]
            if oversized_chunks:
                logger.warning(f"发现 {len(oversized_chunks)} 个分块超过设定大小 {self.chunk_size}")
                logger.warning(f"超大分块大小: {oversized_chunks}")
        
        # 向量化分块
        logger.info(f"开始向量化分块: {self.split_dir}")
        self.vectorize_chunks(chunk_files, delete_existing=delete_existing)
        logger.info(f"向量化完成，集合名称: {self.original_name}")
        
        # 返回处理结果
        return {
            'file_path': file_path,
            'clean_file_path': clean_file_path,
            'metadata_file_path': metadata_file_path,
            'chunk_files': chunk_files,
            'chunk_count': len(chunk_files),
            'chunk_sizes': chunk_sizes,
            'encoding': 'utf-8',
            'collection_name': self.original_name,  # 返回原始集合名称
            'normalized_collection_name': self.normalized_name  # 同时返回规范化后的名称
        }

    def process_directory(self, input_dir: str, extensions: List[str] = None, delete_existing: bool = False, single_thread: bool = False) -> List[Dict[str, Any]]:
        """处理目录中的所有文件
        
        Args:
            input_dir: 输入目录
            extensions: 文件扩展名列表
            delete_existing: 是否删除已存在的集合
            single_thread: 是否使用单线程处理
            
        Returns:
            处理结果列表
        """
        # 使用默认扩展名
        if extensions is None:
            extensions = ['.txt', '.md', '.pdf', '.docx', '.html', '.csv', '.xlsx', '.xls']
        
        logger.info(f"开始处理目录: {input_dir}")
        logger.info(f"支持的文件类型: {', '.join(extensions)}")
        logger.info(f"使用{'单线程' if single_thread else '多线程'}处理")
        
        # 检查集合是否存在
        if delete_existing and self.collection_exists(self.collection_name):
            logger.info(f"集合已存在: {self.original_name}，正在删除...")
            success = self.delete_collection(self.collection_name)
            if not success:
                raise RuntimeError(f"删除集合失败: {self.original_name}")
            logger.info(f"集合已删除: {self.original_name}")
        
        # 获取所有文件
        files = []
        for ext in extensions:
            found_files = list(Path(input_dir).glob(f"**/*{ext}"))
            logger.info(f"找到 {len(found_files)} 个 {ext} 文件")
            files.extend(found_files)
        
        if not files:
            logger.warning(f"目录中没有匹配的文件: {input_dir}")
            return []
        
        logger.info(f"总共找到 {len(files)} 个文件")
        
        # 处理所有文件
        results = []
        failed_files = []
        
        if single_thread:
            # 单线程处理
            for file_path in files:
                try:
                    logger.info(f"开始处理文件: {file_path}")
                    result = self.process_file(str(file_path))
                    results.append(result)
                    logger.info(f"成功处理文件: {file_path}")
                except Exception as e:
                    logger.error(f"处理文件失败: {file_path}, 错误类型: {type(e).__name__}, 错误信息: {str(e)}")
                    failed_files.append((str(file_path), str(e)))
        else:
            # 多线程处理
            with ThreadPoolExecutor(max_workers=self.max_workers) as executor:
                future_to_file = {executor.submit(self.process_file, str(file_path)): file_path for file_path in files}
                
                for future in as_completed(future_to_file):
                    file_path = future_to_file[future]
                    try:
                        result = future.result()
                        results.append(result)
                        logger.info(f"成功处理文件: {file_path}")
                    except Exception as e:
                        logger.error(f"处理文件失败: {file_path}, 错误类型: {type(e).__name__}, 错误信息: {str(e)}")
                        failed_files.append((str(file_path), str(e)))
        
        # 处理结果统计
        success_count = len(results)
        failed_count = len(failed_files)
        total_count = len(files)
        
        logger.info(f"目录处理完成: {input_dir}")
        logger.info(f"总文件数: {total_count}")
        logger.info(f"成功处理: {success_count}")
        logger.info(f"处理失败: {failed_count}")
        
        if failed_files:
            logger.warning("处理失败的文件列表:")
            for file_path, error in failed_files:
                logger.warning(f"- {file_path}: {error}")
        
        # 向量化所有分块
        if results:
            logger.info(f"开始向量化所有分块: {self.split_dir}")
            try:
                self.vectorize_chunks(
                    [file for result in results for file in result['chunk_files']],
                    delete_existing=delete_existing
                )
                logger.info(f"向量化完成，集合名称: {self.original_name}")
            except Exception as e:
                logger.error(f"向量化失败: {str(e)}")
                raise RuntimeError(f"向量化失败: {str(e)}") from e
        
        return results
    
    def process_files(self, file_paths: List[str], delete_existing: bool = False, single_thread: bool = False) -> List[Dict[str, Any]]:
        """处理多个文件
        
        Args:
            file_paths: 文件路径列表
            delete_existing: 是否删除已存在的集合
            single_thread: 是否使用单线程处理
            
        Returns:
            处理结果列表
        """
        logger.info(f"开始处理文件列表，共 {len(file_paths)} 个文件")
        logger.info(f"使用{'单线程' if single_thread else '多线程'}处理")
        
        # 检查集合是否存在
        if delete_existing and self.collection_exists(self.collection_name):
            logger.info(f"集合已存在: {self.original_name}，正在删除...")
            success = self.delete_collection(self.collection_name)
            if not success:
                raise RuntimeError(f"删除集合失败: {self.original_name}")
            logger.info(f"集合已删除: {self.original_name}")
        
        # 按文件类型分组
        file_groups = {}
        for file_path in file_paths:
            _, ext = os.path.splitext(file_path)
            ext = ext.lower()
            if ext not in file_groups:
                file_groups[ext] = []
            file_groups[ext].append(file_path)
        
        # 记录文件类型统计
        for ext, paths in file_groups.items():
            logger.info(f"文件类型 {ext}: {len(paths)} 个文件")
        
        # 处理所有文件
        all_chunks = []  # 存储所有文件的分块结果
        results = []
        failed_files = []
        
        # 按文件类型分别处理
        for ext, group_files in file_groups.items():
            logger.info(f"开始处理 {ext} 类型文件")
            
            if single_thread:
                # 单线程处理
                for file_path in group_files:
                    try:
                        logger.info(f"开始处理文件: {file_path}")
                        # 清洗和分块，但不立即向量化
                        logger.info(f"开始处理文件!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!: {file_path}")
                        result = self._process_single_file(file_path, skip_embedding=True)
                        results.append(result)
                        # 收集分块结果
                        if 'chunk_files' in result:
                            all_chunks.extend(result['chunk_files'])
                        logger.info(f"成功处理文件: {file_path}")
                    except Exception as e:
                        logger.error(f"处理文件失败: {file_path}, 错误类型: {type(e).__name__}, 错误信息: {str(e)}")
                        failed_files.append((file_path, str(e)))
            else:
                # 多线程处理
                with ThreadPoolExecutor(max_workers=self.max_workers) as executor:
                    future_to_file = {
                        executor.submit(self._process_single_file, file_path, skip_embedding=True): file_path 
                        for file_path in group_files
                    }
                    
                    for future in as_completed(future_to_file):
                        file_path = future_to_file[future]
                        try:
                            result = future.result()
                            results.append(result)
                            # 收集分块结果
                            if 'chunk_files' in result:
                                all_chunks.extend(result['chunk_files'])
                            logger.info(f"成功处理文件: {file_path}")
                        except Exception as e:
                            logger.error(f"处理文件失败: {file_path}, 错误类型: {type(e).__name__}, 错误信息: {str(e)}")
                            failed_files.append((file_path, str(e)))
        
        # 处理结果统计
        success_count = len(results)
        failed_count = len(failed_files)
        total_count = len(file_paths)
        
        logger.info("文件列表处理完成")
        logger.info(f"总文件数: {total_count}")
        logger.info(f"成功处理: {success_count}")
        logger.info(f"处理失败: {failed_count}")
        
        if failed_files:
            logger.warning("处理失败的文件列表:")
            for file_path, error in failed_files:
                logger.warning(f"- {file_path}: {error}")
        
        # 统一进行向量化
        if results:
            logger.info(f"开始向量化所有分块，共 {len(all_chunks)} 个分块")
            try:
                logger.info(f"开始向量化所有分块，共++++++++++++++++++++++++++++++++ {len(all_chunks)} 个分块")
                self.vectorize_chunks(all_chunks, delete_existing=delete_existing)
                logger.info(f"向量化完成，集合名称: {self.original_name}")
            except Exception as e:
                logger.error(f"向量化失败: {str(e)}")
                raise RuntimeError(f"向量化失败: {str(e)}") from e
        
        return results

    def _process_single_file(self, file_path: str, skip_embedding: bool = False, delete_existing: bool = False) -> Dict[str, Any]:
        """处理单个文件（清洗和分块，可选是否立即向量化）
        
        Args:
            file_path: 文件路径
            skip_embedding: 是否跳过向量化步骤
            delete_existing: 是否删除现有集合
            
        Returns:
            处理结果
        """
        logger.info(f"开始处理文件: {file_path}")
        
        # 获取文件扩展名
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()
        
        # 根据文件类型选择合适的提取器
        try:
            if ext == '.csv':
                clean_result = self.csv_extractor.extract(file_path)
            elif ext == '.docx':
                word_extractor = WordExtractor(file_path)
                documents = word_extractor.extract()
                if not documents:
                    raise RuntimeError(f"Word文档未提取到任何内容: {file_path}")
                
                # 合并所有文档内容
                text = "\n\n".join(doc.page_content for doc in documents)
                metadata = documents[0].metadata
                clean_result = {
                    "text": text,
                    "metadata": metadata
                }
                logger.info(f"成功从Word文档中提取了 {len(text)} 个字符")
            elif ext == '.pdf':
                pdf_extractor = PdfExtractor(file_path)
                documents = pdf_extractor.extract()
                if not documents:
                    raise RuntimeError(f"PDF文件未提取到任何内容: {file_path}")
                
                # 合并所有页面的内容
                text = "\n\n".join(doc.page_content for doc in documents)
                metadata = {
                    "source": file_path,
                    "total_pages": len(documents),
                    "file_type": "pdf"
                }
                clean_result = {
                    "text": text,
                    "metadata": metadata
                }
                logger.info(f"成功从PDF文件中提取了 {len(documents)} 页，共 {len(text)} 个字符")
            else:
                # 对于其他文本文件，使用通用文本提取器
                clean_result = self.text_extractor.extract(file_path)
                logger.info(f"使用通用文本提取器处理文件: {file_path}")
        except Exception as e:
            logger.error(f"文件提取失败: {file_path}, 错误类型: {type(e).__name__}, 错误信息: {str(e)}")
            raise RuntimeError(f"文件提取失败: {file_path}, 错误: {str(e)}") from e
        
        # 获取清洗后的文本和元数据
        text = clean_result['text']
        metadata = clean_result['metadata']
        
        # 在元数据中添加集合信息
        metadata["collection"] = self.original_name
        metadata["normalized_collection"] = self.normalized_name
        
        # 构建清洗后的文件路径
        clean_file_name = os.path.basename(file_path)
        clean_file_path = os.path.join(self.dataset_dir, clean_file_name)
        
        # 保存清洗后的文本（使用UTF-8编码）
        with open(clean_file_path, 'w', encoding='utf-8') as f:
            f.write(text)
        
        # 保存元数据（使用UTF-8编码）
        metadata_file_path = clean_file_path.replace(ext, '.json')
        with open(metadata_file_path, 'w', encoding='utf-8') as f:
            json.dump(metadata, f, ensure_ascii=False, indent=2)
        
        logger.info(f"文件清洗完成，保存到: {clean_file_path}")
        
        # 分割文本
        chunk_files = self.text_splitter.split_file(
            clean_file_path, 
            output_dir=self.split_dir,
            split_method=self.split_method
        )
        
        # 记录分块信息
        chunk_sizes = []
        for chunk_file in chunk_files:
            try:
                with open(chunk_file, 'r', encoding='utf-8') as f:
                    chunk_text = f.read()
                    chunk_sizes.append(len(chunk_text))
            except Exception as e:
                logger.error(f"读取分块文件失败: {chunk_file}, 错误: {e}")
        
        if chunk_sizes:
            avg_chunk_size = sum(chunk_sizes) / len(chunk_sizes)
            max_chunk_size = max(chunk_sizes)
            min_chunk_size = min(chunk_sizes)
            logger.info(f"分块完成，共 {len(chunk_files)} 个分块")
            logger.info(f"分块大小统计 - 平均: {avg_chunk_size:.2f}, 最大: {max_chunk_size}, 最小: {min_chunk_size}")
            
            # 检查是否有超过设定大小的分块
            oversized_chunks = [size for size in chunk_sizes if size > self.chunk_size]
            if oversized_chunks:
                logger.warning(f"发现 {len(oversized_chunks)} 个分块超过设定大小 {self.chunk_size}")
                logger.warning(f"超大分块大小: {oversized_chunks}")
        
        # 如果不跳过向量化步骤，则立即向量化
        if not skip_embedding and chunk_files:
            logger.info(f"开始向量化分块: {self.split_dir}")
            try:
                logger.info(f"开始向量化分块，集合名称@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@: {self.original_name}")
                self.vectorize_chunks(chunk_files, delete_existing=delete_existing)
                logger.info(f"向量化完成，集合名称: {self.original_name}")
            except Exception as e:
                logger.error(f"向量化失败: {str(e)}")
                raise RuntimeError(f"向量化失败: {str(e)}") from e
        
        return {
            'file_path': file_path,
            'clean_file_path': clean_file_path,
            'metadata_file_path': metadata_file_path,
            'chunk_files': chunk_files,
            'chunk_count': len(chunk_files),
            'chunk_sizes': chunk_sizes,
            'encoding': 'utf-8',
            'collection_name': self.original_name,
            'normalized_collection_name': self.normalized_name
        }
    
    def retrieve(self, query: str, top_k: int = None) -> Dict[str, Any]:
        """检索
        
        Args:
            query: 查询文本
            top_k: 返回结果数量
            
        Returns:
            检索结果
        """
        return self.retriever.retrieve(query, self.collection_name, top_k)
    
    def collection_exists(self, collection_name: str) -> bool:
        """检查集合是否存在
        
        Args:
            collection_name: 集合名称
            
        Returns:
            是否存在
        """
        collections = self.retriever.list_collections()
        return collection_name in collections
    
    def delete_collection(self, collection_name: str) -> bool:
        """删除集合
        
        Args:
            collection_name: 集合名称
            
        Returns:
            是否成功
        """
        try:
            # 获取集合特定的数据库路径
            db_path = os.path.join(config.VECTOR_DB_PATH, collection_name)
            
            # 如果目录存在，则删除
            if os.path.exists(db_path):
                shutil.rmtree(db_path)
                logger.info(f"已删除集合数据库目录: {db_path}")
            
            return True
        except Exception as e:
            logger.error(f"删除集合失败: {collection_name}, 错误: {e}")
            return False

def main():
    """主函数"""
    parser = argparse.ArgumentParser(description='文档处理器')
    parser.add_argument('--action', type=str, required=True, choices=['process_file', 'process_directory', 'process_files', 'retrieve', 'delete_collection', 'list_collections', 'embed_chunks'], help='操作类型')
    parser.add_argument('--input', type=str, help='输入文件或目录')
    parser.add_argument('--inputs', type=str, nargs='+', help='输入文件列表')
    parser.add_argument('--query', type=str, help='查询文本')
    parser.add_argument('--collection_name', type=str, help='集合名称')
    parser.add_argument('--clean_dir', type=str, help='清洗后的文件目录')
    parser.add_argument('--chunk_dir', type=str, help='分块后的文件目录')
    parser.add_argument('--chunk_size', type=int, help='分块大小')
    parser.add_argument('--chunk_overlap', type=int, help='分块重叠大小')
    parser.add_argument('--split_method', type=str, choices=['separator', 'chars', 'tokens', 'markdown'], help='分块方法')
    parser.add_argument('--top_k', type=int, help='返回结果数量')
    parser.add_argument('--extensions', type=str, nargs='+', help='文件扩展名列表')
    parser.add_argument('--delete_existing', action='store_true', help='删除已存在的集合（如果存在）')
    parser.add_argument('--max_workers', type=int, help='最大工作线程数')
    parser.add_argument('--chunks_path', type=str, help='要嵌入的分块文件目录路径')
    
    args = parser.parse_args()
    
    # 检索操作直接使用VectorRetriever，不需要初始化DocumentProcessor
    if args.action == 'retrieve':
        if not args.query:
            parser.error("检索操作需要指定查询文本")
        
        collection_name = args.collection_name or config.COLLECTION_NAME
        logger.info(f"开始检索，集合名称: {collection_name}, 查询: {args.query}")
        
        # 直接使用VectorRetriever进行检索
        retriever = VectorRetriever()
        result = retriever.retrieve(args.query, collection_name, args.top_k)
        
        # 格式化输出
        print(f"查询: {result['query']}")
        print()
        
        for i, (doc, meta, score) in enumerate(zip(result["documents"], result["metadatas"], result["scores"])):
            print(f"结果 {i+1} (分数: {score:.4f}):")
            print(f"来源: {meta.get('source', '未知')}")
            
            # 如果文本太长，则截断
            if len(doc) > 500:
                doc = doc[:500] + "..."
            
            print(f"内容: {doc}")
            print()
        
        return
    
    # 列出集合操作也直接使用VectorRetriever
    elif args.action == 'list_collections':
        retriever = VectorRetriever()
        collections = retriever.list_collections()
        
        if collections:
            print("可用集合:")
            for collection in collections:
                print(f"- {collection}")
        else:
            print("数据库中没有集合")
        
        return
    
    # 删除集合操作也可以直接使用VectorRetriever
    elif args.action == 'delete_collection':
        if not args.collection_name:
            parser.error("删除集合操作需要指定集合名称")
        
        # 获取集合特定的数据库路径
        db_path = os.path.join(config.VECTOR_DB_PATH, args.collection_name)
        
        try:
            # 如果目录存在，则删除
            if os.path.exists(db_path):
                shutil.rmtree(db_path)
                logger.info(f"已删除集合数据库目录: {db_path}")
                print(f"成功删除集合: {args.collection_name}")
            else:
                print(f"集合不存在: {args.collection_name}")
        except Exception as e:
            logger.error(f"删除集合失败: {args.collection_name}, 错误: {e}")
            print(f"删除集合失败: {args.collection_name}")
        
        return
    
    # 对于其他操作，创建DocumentProcessor实例
    processor = DocumentProcessor(
        collection_name=args.collection_name,
        clean_root_dir=args.clean_dir,
        chunk_root_dir=args.chunk_dir,
        chunk_size=args.chunk_size,
        chunk_overlap=args.chunk_overlap,
        split_method=args.split_method,
        max_workers=args.max_workers
    )
    
    if args.action == 'process_file':
        if not args.input:
            parser.error("处理文件操作需要指定输入文件")
        
        result = processor.process_file(args.input, delete_existing=args.delete_existing)
        print(json.dumps(result, ensure_ascii=False, indent=2))
    
    elif args.action == 'process_directory':
        if not args.input:
            parser.error("处理目录操作需要指定输入目录")
        
        result = processor.process_directory(args.input, args.extensions, delete_existing=args.delete_existing)
        print(f"处理完成，共处理了 {len(result)} 个文件")
    
    elif args.action == 'process_files':
        if not args.inputs:
            parser.error("处理多个文件操作需要指定输入文件列表")
        
        result = processor.process_files(args.inputs, delete_existing=args.delete_existing)
        print(f"处理完成，共处理了 {len(result)} 个文件")
    
    elif args.action == 'embed_chunks':
        if not args.chunks_path:
            parser.error("嵌入分块操作需要指定分块文件目录路径")
        
        if not args.collection_name:
            parser.error("嵌入分块操作需要指定集合名称")
        
        # 检查集合是否存在
        retriever = VectorRetriever()
        collections = retriever.list_collections()
        if args.collection_name in collections:
            if args.delete_existing:
                logger.info(f"集合已存在: {args.collection_name}，正在删除...")
                # 获取集合特定的数据库路径
                db_path = os.path.join(config.VECTOR_DB_PATH, args.collection_name)
                
                try:
                    # 如果目录存在，则删除
                    if os.path.exists(db_path):
                        shutil.rmtree(db_path)
                        logger.info(f"已删除集合数据库目录: {db_path}")
                    else:
                        logger.warning(f"集合目录不存在: {db_path}")
                except Exception as e:
                    logger.error(f"删除集合失败: {args.collection_name}, 错误: {e}")
                    parser.error(f"删除集合失败: {args.collection_name}")
                
                logger.info(f"集合已删除: {args.collection_name}")
            else:
                parser.error(f"集合已存在: {args.collection_name}，如需覆盖，请使用 --delete_existing")
        
        # 嵌入分块
        embedder = TextEmbedder()
        embedder.embed_chunks_directory(args.chunks_path, args.collection_name)
        print(f"成功嵌入分块: {args.chunks_path} -> {args.collection_name}")

if __name__ == "__main__":
    main() 